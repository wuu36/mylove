#!/usr/bin/env python3
"""
Word 与 PDF 文档对比与修正系统 - 完整版
主程序：对比差异、生成修正后的 Word 文档
"""

import sys
import os
from pathlib import Path
import shutil

from pdf_parser_v3 import parse_certificate_pdf, DocumentContent
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def compare_content(pdf_content: DocumentContent, docx_path: str) -> dict:
    """对比 PDF 和 Word 内容"""
    doc = Document(docx_path)

    differences = {
        'header': [],
        'title': [],
        'customer': [],
        'batches': [],
        'footer': [],
        'tables': []
    }

    # 提取 Word 文本
    word_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    word_text = '\n'.join(word_paras)

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells)
            word_text += '\n' + row_text

    # 检查标题
    if pdf_content.title and pdf_content.title.upper() not in word_text.upper():
        differences['title'].append(f"缺失标题: {pdf_content.title}")

    # 检查客户信息
    for key, value in pdf_content.customer_info:
        # 检查值是否存在（更宽松的匹配）
        if value and value not in word_text:
            # 尝试去掉空格匹配
            clean_value = value.replace(' ', '')
            clean_text = word_text.replace(' ', '')
            if clean_value not in clean_text:
                differences['customer'].append(f"客户信息差异: {key}: {value}")

    # 检查批次
    for i, batch in enumerate(pdf_content.batches, 1):
        # 提取批次号
        batch_match = batch['header'].split(':')[-1].strip()
        if batch_match and batch_match not in word_text:
            differences['batches'].append(f"批次 {i} 标题缺失: {batch_match}")

        # 检查表格数据 - 更宽松的匹配
        for row in batch['table']:
            # 只检查关键字段
            for cell in row:
                if cell and len(cell) > 2 and not cell.isspace():
                    # 数字和文本分开检查
                    if cell.replace('.', '').replace('-', '').replace(' ', '').isalnum():
                        if cell not in word_text:
                            differences['tables'].append(f"数据缺失: {cell}")
                            break

    # 统计
    total_diffs = sum(len(v) for v in differences.values())

    return {
        'differences': differences,
        'total': total_diffs,
        'identical': total_diffs == 0
    }


def generate_report(comparison: dict) -> str:
    """生成对比报告"""
    lines = ["=" * 60, "对比报告", "=" * 60]

    diffs = comparison['differences']

    if comparison['identical']:
        lines.append("\n[OK] 文档内容一致")
    else:
        lines.append(f"\n发现 {comparison['total']} 处差异:\n")

        for category, items in diffs.items():
            if items:
                lines.append(f"[{category.upper()}] {len(items)} 处")
                for item in items[:5]:
                    lines.append(f"  - {item}")
                if len(items) > 5:
                    lines.append(f"  ... 还有 {len(items) - 5} 处")

    return "\n".join(lines)


def create_corrected_word(pdf_content: DocumentContent, output_path: str):
    """根据 PDF 内容生成 Word 文档"""

    doc = Document()

    # 设置页边距
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    def set_font(run, size=10, bold=False, italic=False):
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    def remove_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    # 页眉表格
    header_table = doc.add_table(rows=6, cols=2)
    remove_borders(header_table)

    header_texts = [
        "Pulcra Chemicals GmbH",
        "Isardamm 79-83",
        "82538 Geretsried",
        "DEUTSCHLAND",
        "",
        pdf_content.date or ""
    ]

    for i, text in enumerate(header_texts):
        cell = header_table.rows[i].cells[1]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                set_font(run, 10)

    doc.add_paragraph()

    # 标题
    if pdf_content.title:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(pdf_content.title)
        set_font(run, 14, bold=True)

    doc.add_paragraph()

    # 客户信息
    info_table = doc.add_table(rows=len(pdf_content.customer_info), cols=2)
    remove_borders(info_table)

    for i, (key, value) in enumerate(pdf_content.customer_info):
        row = info_table.rows[i]
        row.cells[0].text = key + " :"
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, 10)

    doc.add_paragraph()

    # 批次
    for batch in pdf_content.batches:
        para = doc.add_paragraph()
        run = para.add_run(batch['header'])
        set_font(run, 10, bold=True)

        for info in batch['info']:
            para = doc.add_paragraph()
            run = para.add_run(info)
            set_font(run, 10)

        if batch['table']:
            table = doc.add_table(rows=len(batch['table']), cols=5)
            remove_borders(table)

            for r, row_data in enumerate(batch['table']):
                for c, cell_text in enumerate(row_data):
                    if c < len(table.rows[r].cells):
                        cell = table.rows[r].cells[c]
                        cell.text = cell_text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                set_font(run, 9, bold=(r == 0))

            doc.add_paragraph()

    # Released by
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("Released by: ")
    set_font(run, 10, bold=True)
    run = para.add_run("SILKE STEIER")
    set_font(run, 10)

    # Disclaimer
    para = doc.add_paragraph()
    run = para.add_run(
        "The above data represent the results of our quality assessment. "
        "They do not free the purchaser from his own quality check nor do they confirm "
        "that the product has certain properties or is suitable for a specific application."
    )
    set_font(run, 8, italic=True)

    # 页脚
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("DIN EN 10204")
    set_font(run, 10)

    para = doc.add_paragraph()
    run = para.add_run("This certificate is printed out electronically, therefore it has no signature.")
    set_font(run, 7, italic=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Quality Control Department")
    set_font(run, 10)

    doc.save(output_path)


def main():
    # 文件路径
    pdf_path = r"C:\000_claude\mylove\OC-D P241210003.pdf"
    docx_path = r"C:\000_claude\mylove\OC-D P241210003_v8.docx"
    output_path = r"C:\000_claude\mylove\OC-D P241210003_corrected.docx"

    # 命令行参数
    if len(sys.argv) >= 3:
        pdf_path = sys.argv[1]
        docx_path = sys.argv[2]
        output_path = docx_path.replace(".docx", "_corrected.docx")

    print("=" * 60)
    print("Word 与 PDF 文档对比与修正系统")
    print("=" * 60)
    print(f"PDF: {pdf_path}")
    print(f"Word: {docx_path}")
    print()

    # 解析 PDF
    print("1. 解析 PDF...")
    pdf_content = parse_certificate_pdf(pdf_path)
    print(f"   找到 {len(pdf_content.batches)} 个批次")

    # 对比
    print("\n2. 对比文档...")
    comparison = compare_content(pdf_content, docx_path)
    print(generate_report(comparison))

    # 生成修正后的文档
    print("\n3. 生成修正后的 Word 文档...")
    create_corrected_word(pdf_content, output_path)
    print(f"   已保存: {output_path}")

    # 再次对比验证
    print("\n4. 验证修正结果...")
    verification = compare_content(pdf_content, output_path)
    print(generate_report(verification))

    print("\n" + "=" * 60)
    if verification['identical']:
        print("[OK] 修正完成，文档内容一致！")
    else:
        print("仍有差异，可能需要人工检查")
    print("=" * 60)


if __name__ == "__main__":
    main()