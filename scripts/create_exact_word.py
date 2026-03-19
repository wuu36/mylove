#!/usr/bin/env python3
"""
精确复制 PDF 布局到 Word - v37 版本
主要改进:
1. 添加Logo图片支持（从PDF提取）
2. 页眉使用表格布局实现Logo和公司信息左右分布
3. 所有字体改为 TimesNewRoman
4. 使用Tab定位实现精确布局
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import fitz
import os


def set_run_font(run, font_name: str, size_pt: float, bold: bool = False, italic: bool = False):
    """设置 run 的字体属性"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def add_floating_picture(doc, logo_path: str, x_pt: float, y_pt: float, width_pt: float, height_pt: float):
    """添加浮动图片到文档，精确控制位置

    Args:
        doc: Document对象
        logo_path: 图片路径
        x_pt: X位置（点）
        y_pt: Y位置（点）
        width_pt: 图片宽度（点）
        height_pt: 图片高度（点）
    """
    # 添加一个空段落作为锚点
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    # 创建绘图元素
    run = para.add_run()

    # 使用内联图片但通过XML调整位置
    # 添加图片到run
    picture = run.add_picture(logo_path, width=Pt(width_pt), height=Pt(height_pt))

    # 获取drawing元素
    drawing = run._r.find('.//' + qn('w:drawing'))
    if drawing is not None:
        # 设置为浮动图片
        inline = drawing.find(qn('wp:inline'))
        if inline is not None:
            # 将inline转换为anchor（浮动）
            # 创建anchor元素
            anchor = OxmlElement('wp:anchor')
            # 复制inline的属性
            for child in list(inline):
                anchor.append(child)

            # 设置位置属性
            # distT, distB, distL, distR - 文字环绕距离
            anchor.set(qn('wp:distT'), '0')
            anchor.set(qn('wp:distB'), '0')
            anchor.set(qn('wp:distL'), '0')
            anchor.set(qn('wp:distR'), '0')
            anchor.set(qn('wp:simplePos'), '0')
            anchor.set(qn('wp:relativeHeight'), '1')
            anchor.set(qn('wp:behindDoc'), '0')
            anchor.set(qn('wp:locked'), '0')
            anchor.set(qn('wp:layoutInCell'), '1')
            anchor.set(qn('wp:allowOverlap'), '1')

            # 简单位置
            simple_pos = OxmlElement('wp:simplePos')
            simple_pos.set('x', '0')
            simple_pos.set('y', '0')
            anchor.append(simple_pos)

            # 水平位置 - 相对于页面左边缘
            pos_h = OxmlElement('wp:positionH')
            pos_h.set('relativeFrom', 'page')
            pos_offset = OxmlElement('wp:posOffset')
            pos_offset.text = str(int(x_pt * 12700))  # 1pt = 12700 EMU
            pos_h.append(pos_offset)
            anchor.append(pos_h)

            # 垂直位置 - 相对于页面上边缘
            # 由于Y坐标从页面顶部开始，需要减去页边距
            # 页边距=28pt，Logo原始Y=56.8pt
            pos_v = OxmlElement('wp:positionV')
            pos_v.set('relativeFrom', 'page')
            pos_offset = OxmlElement('wp:posOffset')
            # Y = 原始Y坐标（从页面顶部）
            pos_offset.text = str(int(y_pt * 12700))
            pos_v.append(pos_offset)
            anchor.append(pos_v)

            # 替换inline为anchor
            drawing.remove(inline)
            drawing.append(anchor)


def remove_table_borders(table):
    """移除整个表格的边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def pt_to_cm(pt: float) -> float:
    """将点转换为厘米"""
    return pt * 2.54 / 72


def extract_logo_from_pdf(pdf_path: str, output_path: str) -> str:
    """从PDF提取logo图片并保存"""
    doc = fitz.open(pdf_path)
    page = doc[0]
    images = page.get_images()
    if not images:
        doc.close()
        return None

    xref = images[0][0]
    base_image = doc.extract_image(xref)
    image_bytes = base_image["image"]
    image_ext = base_image["ext"]

    logo_path = f"{output_path}.{image_ext}"
    with open(logo_path, "wb") as f:
        f.write(image_bytes)

    doc.close()
    return logo_path


def set_cell_vertical_alignment(cell, align="top"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def create_exact_word_document(output_path: str, pdf_path: str = None, logo_path: str = None):
    """
    根据分析的 PDF 布局创建精确匹配的 Word 文档

    PDF坐标参考 (A4: 595x842 pt, 21x29.7 cm):
    - Logo: bbox=(28.3, 56.8, 166.6, 86.3), 宽约138pt x 高约30pt
    - 页眉公司信息: X=184pt (6.5cm)
    - Page/Date: X=482pt (17cm)
    - 标题: Y=184pt, X=187pt (居中)
    - 客户信息标签: X=28pt
    - 冒号: X=136pt
    - 值: X=140pt
    """
    # 提取或使用已有的logo
    if pdf_path and not logo_path:
        logo_path = extract_logo_from_pdf(pdf_path, "logo")

    doc = Document()

    # 设置页面为 A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4 宽度
    section.page_height = Cm(29.7)  # A4 高度

    # 设置页边距
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # 设置段落默认格式
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # =====================================================
    # 页眉区域 - Tab定位
    # 公司信息用Tab定位到X=184pt
    # 注: Logo添加会影响整体布局，暂时不添加
    # =====================================================

    company_tab = Cm(5.5)
    right_tab = Cm(16.0)

    # 第1行: Pulcra Chemicals GmbH
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(company_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(46)
    run = para.add_run('\tPulcra Chemicals GmbH')
    set_run_font(run, 'TimesNewRoman', 11)

    # 第2行: Isardamm 79-83 + Page: 1/1
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(company_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run('\tIsardamm 79-83\tPage: 1/1')
    set_run_font(run, 'TimesNewRoman', 11)

    # 第3行: 82538 Geretsried
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(company_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run('\t82538 Geretsried')
    set_run_font(run, 'TimesNewRoman', 10)

    # 第4行: DEUTSCHLAND
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(company_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run('\tDEUTSCHLAND')
    set_run_font(run, 'TimesNewRoman', 10)

    # 第5行: 日期
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(8)
    run = para.add_run('\t2025.07.16')
    set_run_font(run, 'TimesNewRoman', 11)

    # 标题前间距
    # 标题Y=184，日期Y=131，间隔53pt
    # 减去日期行高约12pt，还需41pt

    # =====================================================
    # 标题：CERTIFICATE OF ANALYSIS（居中，加粗，16pt）
    # PDF位置: Y=184, 居中
    # =====================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)  # 从34pt增加到40pt
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("CERTIFICATE OF ANALYSIS")
    set_run_font(run, 'TimesNewRoman', 16, bold=True)

    # 客户信息在Y=208，标题在Y=184，间隔24pt
    # 标题是16pt字体，基线到顶部约13pt
    # 标题在Y=185，需要到Customer Y=208
    # 间隔 = 208 - 185 - 13 (标题高度) = 10pt
    # 但需要考虑空段落的高度
    # 直接在标题后添加客户信息，不需要空段落
    # =====================================================
    # 客户信息
    # PDF布局: 标签 X=28pt, 冒号 X=136pt (冒号起始位置), 值 X=140pt
    # 值文本带前导空格: " Jiangsu..."
    # 文本顺序: 标签 -> tab -> ": " -> 值
    # =====================================================
    label_tab = Cm(3.8)   # 冒号起始位置 (108pt from margin)

    customer_data = [
        ("Customer", "Jiangsu Mingxin Xuteng Technology C"),
        ("Customer Nr.", "1508866"),
        ("Product Name", "FORYL OC-D(I104)"),
        ("Product Nr.", "20660"),
    ]

    # 直接添加客户信息，调整间距
    for i, (label, value) in enumerate(customer_data):
        para = doc.add_paragraph()
        para.paragraph_format.tab_stops.add_tab_stop(label_tab, WD_TAB_ALIGNMENT.LEFT)
        # 第一个客户信息需要额外间距 (标题Y=185 + 标题高度13pt = 198, 到Y=208还需10pt)
        # 但实际生成Y=214，偏移+6，所以减少space_before
        if i == 0:
            para.paragraph_format.space_before = Pt(5)  # 减少间距
        else:
            para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # 标签加粗
        run1 = para.add_run(label)
        set_run_font(run1, 'TimesNewRoman', 11, bold=True)
        # Tab + 冒号 + 空格 (加粗)
        run2 = para.add_run('\t: ')
        set_run_font(run2, 'TimesNewRoman', 11, bold=True)
        # 值
        run3 = para.add_run(value)
        set_run_font(run3, 'TimesNewRoman', 11)

    # =====================================================
    # 批次信息
    # =====================================================
    add_batch_section(doc, "P241200117", "2024.04.29", "2026.04.29", "40000403286",
                      "6.9", "35.3", first_batch=True)
    add_batch_section(doc, "P241210003", "2024.04.30", "2026.04.30", "90000071881",
                      "6.8", "35.7", first_batch=False)

    # =====================================================
    # Released by
    # Released偏移+39.9pt(位置偏低)，需要减少space_before使其上升
    # =====================================================
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)  # 从20pt减少到10pt
    run = para.add_run("Released by: ")
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run("SILKE STEIER")
    set_run_font(run, 'TimesNewRoman', 11, bold=True)

    # =====================================================
    # 免责声明（正常字体，11pt）
    # 原PDF分成3行，需要保持一致
    # =====================================================
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("The above data represent the results of our quality assessment.")
    set_run_font(run, 'TimesNewRoman', 11, italic=False)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("They do not free the purchaser from his own quality check nor do they confirm that the product has certain properties or")
    set_run_font(run, 'TimesNewRoman', 11, italic=False)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("is suitable for a specific application.")
    set_run_font(run, 'TimesNewRoman', 11, italic=False)

    # =====================================================
    # 页脚
    # =====================================================
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(130)  # v33最佳值
    run = para.add_run("DIN EN 10204")
    set_run_font(run, 'TimesNewRoman', 8)

    para = doc.add_paragraph()
    run = para.add_run("This certificate is printed out electronically, therefore it has no signature.")
    set_run_font(run, 'TimesNewRoman', 8)

    para = doc.add_paragraph()
    run = para.add_run("Quality Control Department")
    set_run_font(run, 'TimesNewRoman', 8, bold=True)

    doc.save(output_path)
    print(f"文档已保存: {output_path}")


def add_batch_section(doc, batch_no: str, prod_date: str, exp_date: str, lot_no: str,
                      ph_result: str, water_result: str, first_batch: bool = False):
    """添加批次信息部分

    表格列宽按PDF X坐标计算 (pt转cm):
    - Specification: X=28pt (1cm), 宽度=201-28=173pt (6.1cm)
    - Method: X=201pt (7.1cm), 宽度=280-201=79pt (2.8cm)
    - Unit: X=280pt (9.9cm), 宽度=316-280=36pt (1.3cm)
    - Result: X=316pt (11.2cm), 宽度=416-316=100pt (3.5cm)
    - Standard: X=416pt (14.7cm), 宽度至边距
    """
    # Tab位置用于冒号对齐 (与客户信息相同)
    label_tab = Cm(3.8)  # 冒号起始位置

    # Batch Number - 第一个批次需要额外间距
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(label_tab, WD_TAB_ALIGNMENT.LEFT)
    if first_batch:
        para.paragraph_format.space_before = Pt(43)  # 从46pt减少到43pt
    else:
        para.paragraph_format.space_before = Pt(18)
    run = para.add_run("Batch Number")
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run('\t: ')
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run(batch_no)
    set_run_font(run, 'TimesNewRoman', 11)

    # Production Date
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(label_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("Production Date")
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run('\t: ')
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run(prod_date)
    set_run_font(run, 'TimesNewRoman', 11)

    # Expiration Date
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(label_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("Expiration Date")
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run('\t: ')
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run(exp_date)
    set_run_font(run, 'TimesNewRoman', 11)

    # Inspection Lot
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(label_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("Inspection Lot")
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run('\t: ')
    set_run_font(run, 'TimesNewRoman', 11, bold=True)
    run = para.add_run(lot_no)
    set_run_font(run, 'TimesNewRoman', 11)

    # 表格区域 - 使用段落+Tab替代表格，精确控制行间距
    # PDF位置: Specification Y=356, 分隔线 Y=365, AUSSEHEN Y=373, PH Y=390, WASSERGEHALT Y=407
    # 表头Tab位置: Method X=201, Unit X=280, Result X=316, Standard X=416
    spec_tab = Cm(6.1)   # 201-28=173pt=6.1cm
    unit_tab = Cm(8.9)   # 280-28=252pt=8.9cm
    result_tab = Cm(10.1) # 316-28=288pt=10.1cm
    std_tab = Cm(13.7)   # 416-28=388pt=13.7cm

    # 表头行 - Specification偏移+8pt，需要减少space_before
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)  # 从10pt减少到2pt
    para.paragraph_format.tab_stops.add_tab_stop(spec_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(unit_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(result_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(std_tab, WD_TAB_ALIGNMENT.LEFT)
    run = para.add_run("Specification\tMethod\tUnit\tResult\tStandard")
    set_run_font(run, 'TimesNewRoman', 8, bold=True)

    # 分隔线行
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.tab_stops.add_tab_stop(spec_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(unit_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(result_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(std_tab, WD_TAB_ALIGNMENT.LEFT)
    # 使用完整的分隔线字符串（与原PDF匹配）
    run = para.add_run("_______________________________ _______________________________ _______________________________ _______________________________ _")
    set_run_font(run, 'TimesNewRoman', 8, bold=True)  # 原始PDF是8pt粗体

    # 数据行 1: AUSSEHEN (Y=373)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.tab_stops.add_tab_stop(spec_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(unit_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(result_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(std_tab, WD_TAB_ALIGNMENT.LEFT)
    run = para.add_run(f"AUSSEHEN;20°C\t\t\tCOLORLESS TO YELLOWI\tCOLORLESS TO YELLOWISH")
    set_run_font(run, 'TimesNewRoman', 8)

    # 数据行 2: PH (Y=390，间隔17pt)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(5)  # 增加行间距
    para.paragraph_format.tab_stops.add_tab_stop(spec_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(unit_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(result_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(std_tab, WD_TAB_ALIGNMENT.LEFT)
    run = para.add_run(f"PH;10%\t\t\t{ph_result}\t6.0 -8.0")
    set_run_font(run, 'TimesNewRoman', 8)

    # 数据行 3: WASSERGEHALT (Y=407，间隔17pt)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(5)  # 增加行间距
    para.paragraph_format.tab_stops.add_tab_stop(spec_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(unit_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(result_tab, WD_TAB_ALIGNMENT.LEFT)
    para.paragraph_format.tab_stops.add_tab_stop(std_tab, WD_TAB_ALIGNMENT.LEFT)
    run = para.add_run(f"WASSERGEHALT,KARL FISCHER\t\t%\t{water_result}\t34.0 -37.0")
    set_run_font(run, 'TimesNewRoman', 8)

    # 批次后空行
    doc.add_paragraph()


if __name__ == "__main__":
    output_path = r"C:\000_claude\mylove\OC-D P241210003_v45.docx"
    pdf_path = r"C:\000_claude\mylove\OC-D P241210003.pdf"
    logo_path = r"C:\000_claude\mylove\logo_correct.png"
    create_exact_word_document(output_path, pdf_path, logo_path)